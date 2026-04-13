"""
מצפן AI — סקירה יומית
Daily AI briefing for Arie Steinfeld
Searches for AI news, tools, competitors, and tips — generates a Word doc and sends via Telegram.
"""

import os
import json
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

import anthropic

# --- Configuration: config.py locally, env vars on GitHub Actions ---
try:
    from config import ANTHROPIC_API_KEY, TELEGRAM_BOT_TOKEN, TELEGRAM_CHAT_ID
except ImportError:
    ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
    TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN")
    TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID")

# --- Configuration ---
_BASE_DIR = Path(__file__).parent
REPORTS_DIR = _BASE_DIR / "reports"
SEEN_URLS_FILE = _BASE_DIR / "seen_urls.json"
MAX_SEEN_URLS = 500  # Keep last 500 URLs to avoid file growing forever


def load_seen_urls():
    """Load list of URLs already shown in previous briefings."""
    if SEEN_URLS_FILE.exists():
        data = json.loads(SEEN_URLS_FILE.read_text(encoding="utf-8"))
        return data.get("urls", [])
    return []


def save_seen_urls(new_urls):
    """Add new URLs to the seen list, keeping only the last MAX_SEEN_URLS."""
    existing = load_seen_urls()
    combined = existing + [u for u in new_urls if u not in existing]
    combined = combined[-MAX_SEEN_URLS:]  # Keep only last N
    SEEN_URLS_FILE.write_text(
        json.dumps({"urls": combined}, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )


def extract_urls(content):
    """Extract all URLs from briefing markdown content."""
    return re.findall(r'\(https?://[^\)]+\)', content)

# --- Search Topics ---
SEARCH_QUERIES = {
    "tools": [
        "AI tools for small business 2026",
        "new AI tools launched this week",
        "Claude AI updates March 2026",
        "AI automation tools no-code 2026",
    ],
    "competitors": [
        "יעוץ AI לעסקים קטנים ישראל 2026",
        "AI consulting small business Israel",
        "AI implementation SMB consulting",
    ],
    "tips": [
        "Claude Code tips and tricks 2026",
        "AI productivity tips for consultants",
        "prompt engineering best practices 2026",
    ],
    "trends": [
        "AI regulation Israel 2026",
        "AI adoption trends small business",
        "AI case studies SMB 2026",
    ],
}


def run_claude_research():
    """Use Anthropic Python SDK with web_search to generate the briefing content."""
    today = datetime.now().strftime("%Y-%m-%d")
    prompt_file = Path(__file__).parent / "briefing_prompt.txt"
    prompt_text = prompt_file.read_text(encoding="utf-8")

    # Build seen URLs block
    seen_urls = load_seen_urls()
    if seen_urls:
        seen_block = (
            "\n\nCRITICAL - PREVIOUSLY SHOWN ARTICLES (DO NOT include these again):\n"
            + "\n".join(f"- {u}" for u in seen_urls[-200:])
            + "\n\nOnly include news published in the last 48 hours. Skip anything older."
        )
    else:
        seen_block = "\n\nOnly include news published in the last 48 hours."

    prompt = f"Today is {today}.\n\n{prompt_text}{seen_block}"

    import time

    for attempt in range(3):
        try:
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

            messages = [{"role": "user", "content": prompt}]
            full_text = ""

            while True:
                response = client.messages.create(
                    model="claude-sonnet-4-5",
                    max_tokens=4096,
                    tools=[{"type": "web_search_20250305", "name": "web_search"}],
                    messages=messages,
                )

                # Collect all text from this response
                for block in response.content:
                    if hasattr(block, "text"):
                        full_text += block.text

                if response.stop_reason == "end_turn":
                    break

                if response.stop_reason == "tool_use":
                    tool_results = []
                    for block in response.content:
                        if block.type == "tool_use":
                            tool_results.append({
                                "type": "tool_result",
                                "tool_use_id": block.id,
                                "content": "",
                            })
                    messages.append({"role": "assistant", "content": response.content})
                    messages.append({"role": "user", "content": tool_results})
                else:
                    break

            return full_text.strip() if full_text.strip() else None

        except (anthropic.RateLimitError, anthropic.APIStatusError) as e:
            # Retry on rate limit (429) or overloaded (529)
            if isinstance(e, anthropic.APIStatusError) and e.status_code not in (429, 529):
                print(f"Anthropic SDK error: {e}")
                return None
            wait = 60 * (attempt + 1)
            print(f"API busy (attempt {attempt + 1}/3), waiting {wait}s... [{e}]")
            time.sleep(wait)
        except Exception as e:
            print(f"Anthropic SDK error: {e}")
            return None

    print("Failed after 3 attempts.")
    return None


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})

    new_run = paragraph._element.makeelement(qn("w:r"), {})
    rPr = paragraph._element.makeelement(qn("w:rPr"), {})

    color = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    rPr.append(color)

    u = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_rich_text(paragraph, text):
    """Parse markdown inline formatting (bold + links) and add to paragraph."""
    # Pattern: split by bold and links
    # Process links first: [text](url)
    link_pattern = r'\[([^\]]+)\]\((https?://[^\)]+)\)'
    bold_pattern = r'\*\*(.*?)\*\*'

    # Combined pattern to split by both links and bold
    combined = r'(\[(?:[^\]]+)\]\((?:https?://[^\)]+)\)|\*\*(?:.*?)\*\*)'
    parts = re.split(combined, text)

    for part in parts:
        if not part:
            continue
        link_match = re.match(link_pattern, part)
        bold_match = re.match(bold_pattern, part)
        if link_match:
            add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
        else:
            paragraph.add_run(part)


def create_word_document(content, filepath):
    """Create an RTL Word document from the briefing content."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Set RTL for the entire document
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style configuration
    style = doc.styles["Normal"]
    font = style.font
    font.name = "David"
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Set RTL on Normal style
    pPr = style.element.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)

    today = datetime.now().strftime("%d/%m/%Y")

    # Title
    title = doc.add_heading(f"🧭 מצפן AI — סקירה יומית — {today}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in title.runs:
        run.font.size = Pt(20)
    # Set RTL on title
    tPr = title._element.get_or_add_pPr()
    tBidi = tPr.makeelement(qn("w:bidi"), {})
    tPr.append(tBidi)

    # Parse markdown content into the document
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("## "):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hPr = heading._element.get_or_add_pPr()
            hBidi = hPr.makeelement(qn("w:bidi"), {})
            hPr.append(hBidi)
        elif line.startswith("### "):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hPr = heading._element.get_or_add_pPr()
            hBidi = hPr.makeelement(qn("w:bidi"), {})
            hPr.append(hBidi)
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_rich_text(para, line[2:])
            pPr = para._element.get_or_add_pPr()
            bidi = pPr.makeelement(qn("w:bidi"), {})
            pPr.append(bidi)
        else:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_rich_text(para, line)
            pPr_elem = para._element.get_or_add_pPr()
            bidi = pPr_elem.makeelement(qn("w:bidi"), {})
            pPr_elem.append(bidi)

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph("—")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = doc.add_paragraph(f"נוצר אוטומטית על ידי מצפן AI — {today}")
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_text.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(str(filepath))
    return filepath


def send_telegram_message(text):
    """Send a text message via Telegram bot."""
    import requests

    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    data = {"chat_id": TELEGRAM_CHAT_ID, "text": text, "parse_mode": "HTML"}
    resp = requests.post(url, json=data)
    return resp.json().get("ok", False)


def send_telegram_document(filepath, caption=""):
    """Send a document via Telegram bot."""
    import requests

    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendDocument"
    with open(filepath, "rb") as f:
        files = {"document": (filepath.name, f)}
        data = {"chat_id": TELEGRAM_CHAT_ID, "caption": caption}
        resp = requests.post(url, data=data, files=files)
    return resp.json().get("ok", False)


def create_telegram_summary(content):
    """Create a short summary for the Telegram message."""
    lines = content.split("\n")
    summary_lines = []
    in_summary = False

    for line in lines:
        if "שורה תחתונה" in line:
            in_summary = True
            continue
        if in_summary and line.strip():
            summary_lines.append(line.strip())
        if in_summary and line.startswith("##"):
            break

    today = datetime.now().strftime("%d/%m/%Y")
    summary = "\n".join(summary_lines) if summary_lines else "הסקירה היומית מוכנה!"

    # Clean markdown
    summary = re.sub(r"\*\*(.*?)\*\*", r"\1", summary)
    summary = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", summary)

    return f"🧭 מצפן AI — {today}\n\n{summary}\n\n📎 הקובץ המלא מצורף"


def main():
    print(f"[{datetime.now()}] Starting AI Daily Briefing...")

    # Step 1: Research
    print("Researching AI news...")
    content = run_claude_research()
    if not content:
        error_msg = "⚠️ מצפן AI: לא הצלחתי ליצור סקירה היום. בדוק את הלוגים."
        send_telegram_message(error_msg)
        print("Failed to generate content.")
        sys.exit(1)

    # Step 2: Create Word document
    today = datetime.now().strftime("%Y-%m-%d")
    filename = f"ai-briefing-{today}.docx"
    filepath = REPORTS_DIR / filename
    print(f"Creating Word document: {filepath}")
    create_word_document(content, filepath)

    # Step 3: Send via Telegram
    print("Sending to Telegram...")
    summary = create_telegram_summary(content)
    send_telegram_message(summary)
    send_telegram_document(filepath, caption=f"מצפן AI — סקירה יומית — {today}")

    # Step 4: Save seen URLs to avoid repetition tomorrow
    new_urls = extract_urls(content)
    save_seen_urls(new_urls)
    print(f"Saved {len(new_urls)} URLs to seen list.")

    print(f"[{datetime.now()}] Done! Report saved to {filepath}")


if __name__ == "__main__":
    main()
